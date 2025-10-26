"""
Streaming Document Processor for Large Files
Handles documents of any size using batch processing and memory management
"""

import gc
import psutil
import hashlib
from typing import Dict, List, Any, Generator, Optional, Callable
from datetime import datetime
import PyPDF2
import docx
import io
from advanced_chunking import LegalSemanticChunker, extract_pdf_text, extract_docx_text


class MemoryMonitor:
    """Monitor and manage memory usage during processing"""

    @staticmethod
    def get_memory_usage() -> Dict[str, float]:
        """Get current memory usage statistics"""
        process = psutil.Process()
        memory_info = process.memory_info()

        return {
            'used_mb': memory_info.rss / 1024 / 1024,
            'percent': process.memory_percent(),
            'available_mb': psutil.virtual_memory().available / 1024 / 1024
        }

    @staticmethod
    def cleanup():
        """Force garbage collection and cleanup"""
        gc.collect()
        # Clear any torch cache if available
        try:
            import torch
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
        except ImportError:
            pass

    @staticmethod
    def get_memory_status() -> str:
        """Get memory status as a color indicator"""
        mem = MemoryMonitor.get_memory_usage()
        percent = mem['percent']

        if percent < 60:
            return 'green'
        elif percent < 80:
            return 'yellow'
        else:
            return 'red'


class StreamingDocumentProcessor:
    """Process large documents in streaming batches to prevent memory overflow"""

    def __init__(
        self,
        chunker: LegalSemanticChunker,
        embedding_model,
        chunk_batch_size: int = 500,
        embedding_batch_size: int = 500,
        upload_batch_size: int = 100
    ):
        self.chunker = chunker
        self.embedding_model = embedding_model
        self.chunk_batch_size = chunk_batch_size
        self.embedding_batch_size = embedding_batch_size
        self.upload_batch_size = upload_batch_size
        self.memory_monitor = MemoryMonitor()

    def extract_text_streaming(
        self,
        file_obj,
        file_type: str,
        progress_callback: Optional[Callable] = None
    ) -> Generator[str, None, None]:
        """Extract text from document in chunks for streaming processing"""

        if file_type == "application/pdf":
            yield from self._extract_pdf_streaming(file_obj, progress_callback)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            yield from self._extract_docx_streaming(file_obj, progress_callback)
        else:
            # For text files, read in chunks
            yield from self._extract_text_streaming(file_obj, progress_callback)

    def _extract_pdf_streaming(
        self,
        file_obj,
        progress_callback: Optional[Callable] = None
    ) -> Generator[str, None, None]:
        """Extract PDF text page by page"""
        try:
            file_obj.seek(0)
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_obj.read()))
            total_pages = len(pdf_reader.pages)

            # Process pages in batches
            batch_size = 50
            for batch_start in range(0, total_pages, batch_size):
                batch_end = min(batch_start + batch_size, total_pages)
                batch_text = ""

                for page_num in range(batch_start, batch_end):
                    page = pdf_reader.pages[page_num]
                    batch_text += page.extract_text() + "\n"

                    if progress_callback:
                        progress_callback({
                            'phase': 'text_extraction',
                            'current': page_num + 1,
                            'total': total_pages,
                            'message': f'Extracting page {page_num + 1} of {total_pages}'
                        })

                yield batch_text

                # Cleanup memory after each batch
                self.memory_monitor.cleanup()

        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")

    def _extract_docx_streaming(
        self,
        file_obj,
        progress_callback: Optional[Callable] = None
    ) -> Generator[str, None, None]:
        """Extract DOCX text in paragraph batches"""
        try:
            file_obj.seek(0)
            doc = docx.Document(io.BytesIO(file_obj.read()))
            total_paragraphs = len(doc.paragraphs)

            # Process paragraphs in batches
            batch_size = 100
            for batch_start in range(0, total_paragraphs, batch_size):
                batch_end = min(batch_start + batch_size, total_paragraphs)
                batch_text = ""

                for i in range(batch_start, batch_end):
                    batch_text += doc.paragraphs[i].text + "\n"

                    if progress_callback:
                        progress_callback({
                            'phase': 'text_extraction',
                            'current': i + 1,
                            'total': total_paragraphs,
                            'message': f'Extracting paragraph {i + 1} of {total_paragraphs}'
                        })

                yield batch_text

                # Cleanup memory after each batch
                self.memory_monitor.cleanup()

        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")

    def _extract_text_streaming(
        self,
        file_obj,
        progress_callback: Optional[Callable] = None
    ) -> Generator[str, None, None]:
        """Extract plain text in chunks"""
        try:
            file_obj.seek(0)
            chunk_size = 1024 * 1024  # 1MB chunks

            while True:
                chunk = file_obj.read(chunk_size)
                if not chunk:
                    break

                text_chunk = chunk.decode('utf-8', errors='ignore')
                yield text_chunk

                if progress_callback:
                    progress_callback({
                        'phase': 'text_extraction',
                        'message': f'Reading text file...'
                    })

        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")

    def process_document_streaming(
        self,
        file_obj,
        file_name: str,
        file_type: str,
        content_hash: str,
        progress_callback: Optional[Callable] = None,
        checkpoint_callback: Optional[Callable] = None
    ) -> Generator[Dict[str, Any], None, None]:
        """
        Process document in streaming batches
        Yields batches of processed chunks ready for embedding and upload
        """

        # Phase 1: Extract text in streaming fashion
        if progress_callback:
            progress_callback({
                'phase': 'text_extraction',
                'message': 'Starting text extraction...'
            })

        accumulated_text = ""
        total_text_length = 0

        for text_batch in self.extract_text_streaming(file_obj, file_type, progress_callback):
            accumulated_text += text_batch
            total_text_length += len(text_batch)

        if progress_callback:
            progress_callback({
                'phase': 'text_extraction',
                'completed': True,
                'total_characters': total_text_length,
                'message': f'Extracted {total_text_length:,} characters'
            })

        # Phase 2: Chunk the document
        if progress_callback:
            progress_callback({
                'phase': 'chunking',
                'message': 'Starting semantic chunking...'
            })

        # Create chunks from the accumulated text
        chunks = self.chunker.legal_aware_chunking(accumulated_text, max_chunk_size=1200)

        if not chunks:
            raise Exception("No chunks were created - check file format")

        total_chunks = len(chunks)

        if progress_callback:
            progress_callback({
                'phase': 'chunking',
                'completed': True,
                'total_chunks': total_chunks,
                'message': f'Created {total_chunks:,} semantic chunks'
            })

        # Clear the accumulated text to free memory
        del accumulated_text
        self.memory_monitor.cleanup()

        # Phase 3: Process chunks in batches
        batch_num = 0
        total_batches = (total_chunks + self.chunk_batch_size - 1) // self.chunk_batch_size

        for i in range(0, total_chunks, self.chunk_batch_size):
            batch_num += 1
            batch_chunks = chunks[i:i + self.chunk_batch_size]

            if progress_callback:
                progress_callback({
                    'phase': 'processing',
                    'batch_num': batch_num,
                    'total_batches': total_batches,
                    'chunks_in_batch': len(batch_chunks),
                    'chunks_processed': i,
                    'total_chunks': total_chunks,
                    'message': f'Processing batch {batch_num} of {total_batches}'
                })

            # Yield this batch for embedding and upload
            yield {
                'batch_num': batch_num,
                'total_batches': total_batches,
                'chunks': batch_chunks,
                'chunks_processed': i,
                'total_chunks': total_chunks,
                'file_name': file_name,
                'content_hash': content_hash
            }

            # Create checkpoint after each batch
            if checkpoint_callback:
                checkpoint_callback({
                    'phase': 'processing',
                    'batch_num': batch_num,
                    'chunks_processed': i + len(batch_chunks),
                    'total_chunks': total_chunks,
                    'memory_usage': self.memory_monitor.get_memory_usage()
                })

            # Cleanup memory after batch
            self.memory_monitor.cleanup()

    def generate_embeddings_batch(
        self,
        chunks: List[Dict],
        progress_callback: Optional[Callable] = None
    ) -> List[Any]:
        """Generate embeddings for a batch of chunks"""

        chunk_texts = [ch['text'] for ch in chunks]
        total_chunks = len(chunk_texts)

        # Process in smaller sub-batches for embedding generation
        all_embeddings = []
        sub_batch_size = self.embedding_batch_size
        total_sub_batches = (total_chunks + sub_batch_size - 1) // sub_batch_size

        for i in range(0, total_chunks, sub_batch_size):
            sub_batch = chunk_texts[i:i + sub_batch_size]

            if progress_callback:
                progress_callback({
                    'phase': 'embedding',
                    'current': i + len(sub_batch),
                    'total': total_chunks,
                    'sub_batch': (i // sub_batch_size) + 1,
                    'total_sub_batches': total_sub_batches,
                    'message': f'Generating embeddings {i + 1}-{i + len(sub_batch)} of {total_chunks}'
                })

            # Generate embeddings for this sub-batch
            embeddings = self.embedding_model.encode(sub_batch, show_progress_bar=False)
            all_embeddings.extend(embeddings)

            # Cleanup after each sub-batch
            self.memory_monitor.cleanup()

        return all_embeddings
